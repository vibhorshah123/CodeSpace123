"""
Script to generate a comprehensive How-To Word document for D365 Python Comparison Tool
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx package...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_bullet_point(doc, text, level=0):
    """Add a bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    return para

def add_numbered_list(doc, text):
    """Add a numbered list item"""
    return doc.add_paragraph(text, style='List Number')

def add_code_block(doc, code_text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Light gray background simulation
    return para

def create_how_to_document():
    """Create the comprehensive How-To Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========================================================================
    # TITLE PAGE
    # ========================================================================
    title = doc.add_heading('Dynamics 365 Environment Comparison Tool', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('How-To Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    
    doc.add_paragraph()
    version_para = doc.add_paragraph('Version 1.0')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph('November 2025')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ========================================================================
    # TABLE OF CONTENTS
    # ========================================================================
    add_heading(doc, 'Table of Contents', 1)
    add_bullet_point(doc, 'Overview')
    add_bullet_point(doc, 'Prerequisites')
    add_bullet_point(doc, 'Installation & Setup')
    add_bullet_point(doc, 'Getting Started')
    add_bullet_point(doc, 'Feature Guides')
    add_bullet_point(doc, 'Schema Comparison', level=1)
    add_bullet_point(doc, 'Data Comparison', level=1)
    add_bullet_point(doc, 'Flow Comparison', level=1)
    add_bullet_point(doc, 'Solution Comparison', level=1)
    add_bullet_point(doc, 'Excel Reports')
    add_bullet_point(doc, 'Authentication Methods')
    add_bullet_point(doc, 'Troubleshooting')
    add_bullet_point(doc, 'Tips & Best Practices')
    
    doc.add_page_break()
    
    # ========================================================================
    # OVERVIEW
    # ========================================================================
    add_heading(doc, '1. Overview', 1)
    
    add_paragraph(doc, 'The Dynamics 365 Environment Comparison Tool is a Python-based command-line application that helps you compare configurations and data between different Dynamics 365 environments.')
    
    add_heading(doc, 'Key Features', 2)
    add_bullet_point(doc, 'Schema Comparison: Compare table structures, fields, and properties')
    add_bullet_point(doc, 'Data Comparison: Compare actual records with field-level differences')
    add_bullet_point(doc, 'Flow Comparison: Compare Power Automate flows between environments')
    add_bullet_point(doc, 'Solution Comparison: Compare all solution components (160+ types)')
    add_bullet_point(doc, 'Excel Reports: Generate detailed, color-coded comparison reports')
    add_bullet_point(doc, 'Interactive Authentication: Secure OAuth 2.0 with browser-based login')
    add_bullet_point(doc, 'Relationship Discovery: Automatically compare related records')
    
    add_heading(doc, 'When to Use This Tool', 2)
    add_bullet_point(doc, 'Before deploying solutions between environments')
    add_bullet_point(doc, 'After migrations to validate data integrity')
    add_bullet_point(doc, 'Troubleshooting environment differences')
    add_bullet_point(doc, 'Documentation and audit trails')
    add_bullet_point(doc, 'Quality assurance testing')
    
    doc.add_page_break()
    
    # ========================================================================
    # PREREQUISITES
    # ========================================================================
    add_heading(doc, '2. Prerequisites', 1)
    
    add_heading(doc, 'System Requirements', 2)
    add_bullet_point(doc, 'Operating System: Windows 7 or later')
    add_bullet_point(doc, 'Python: Version 3.8 or higher')
    add_bullet_point(doc, 'Internet Connection: Required for API access')
    add_bullet_point(doc, 'Disk Space: Minimum 100 MB free')
    
    add_heading(doc, 'Access Requirements', 2)
    add_bullet_point(doc, 'Valid Dynamics 365 credentials for both environments')
    add_bullet_point(doc, 'System Administrator or System Customizer role (recommended)')
    add_bullet_point(doc, 'Read permissions on all entities you want to compare')
    add_bullet_point(doc, 'Network access to both D365 environments')
    
    add_heading(doc, 'Python Dependencies', 2)
    add_paragraph(doc, 'The tool automatically installs these when you run setup.bat:')
    add_bullet_point(doc, 'requests >= 2.31.0')
    add_bullet_point(doc, 'openpyxl >= 3.1.2')
    add_bullet_point(doc, 'tqdm >= 4.66.0')
    
    doc.add_page_break()
    
    # ========================================================================
    # INSTALLATION & SETUP
    # ========================================================================
    add_heading(doc, '3. Installation & Setup', 1)
    
    add_heading(doc, 'Step 1: Download or Clone the Repository', 2)
    add_paragraph(doc, 'Navigate to the project folder:')
    add_code_block(doc, 'c:\\Users\\v-vibhorshah\\source\\repos\\CodeSpace123\\D365PythonComparison')
    
    add_heading(doc, 'Step 2: Install Dependencies', 2)
    add_paragraph(doc, 'Option A: Easy Setup (Recommended)', bold=True)
    add_paragraph(doc, 'Simply double-click setup.bat in the project folder. This will automatically install all required Python packages.')
    
    add_paragraph(doc, 'Option B: Manual Setup', bold=True)
    add_paragraph(doc, 'Open Command Prompt or PowerShell and run:')
    add_code_block(doc, 'cd c:\\Users\\v-vibhorshah\\source\\repos\\CodeSpace123\\D365PythonComparison\npip install -r requirements.txt')
    
    add_heading(doc, 'Step 3: Verify Installation', 2)
    add_paragraph(doc, 'Run the tool to verify everything is working:')
    add_code_block(doc, 'python main.py')
    add_paragraph(doc, 'You should see the welcome banner and authentication prompt.')
    
    doc.add_page_break()
    
    # ========================================================================
    # GETTING STARTED
    # ========================================================================
    add_heading(doc, '4. Getting Started', 1)
    
    add_heading(doc, 'Launching the Tool', 2)
    add_paragraph(doc, 'There are three ways to start the tool:')
    
    add_paragraph(doc, 'Method 1: Batch File (Easiest)', bold=True)
    add_paragraph(doc, 'Double-click run.bat in the project folder.')
    
    add_paragraph(doc, 'Method 2: Python Command', bold=True)
    add_code_block(doc, 'python main.py')
    
    add_paragraph(doc, 'Method 3: Standalone Executable (No Python Required)', bold=True)
    add_paragraph(doc, 'First, build the executable:')
    add_code_block(doc, 'build_exe.bat')
    add_paragraph(doc, 'Then run:')
    add_code_block(doc, 'dist\\D365ComparisonTool.exe')
    
    add_heading(doc, 'First-Time Setup Wizard', 2)
    add_paragraph(doc, 'When you launch the tool, you will see:')
    
    add_numbered_list(doc, 'Authentication Method Selection')
    add_paragraph(doc, 'Choose between Interactive Login (Option 1) or Service Principal (Option 2).')
    
    add_numbered_list(doc, 'Environment Configuration')
    add_paragraph(doc, 'Enter just the environment name (e.g., "mashtest" or "orgname").')
    add_paragraph(doc, 'The tool automatically constructs the full URL: https://mashtest.crm.dynamics.com')
    
    add_numbered_list(doc, 'Comparison Type Selection')
    add_paragraph(doc, 'Choose from the main menu:')
    add_bullet_point(doc, '1. Schema Comparison', level=1)
    add_bullet_point(doc, '2. Data Comparison', level=1)
    add_bullet_point(doc, '3. Flow Comparison', level=1)
    add_bullet_point(doc, '4. Solution Comparison', level=1)
    
    doc.add_page_break()
    
    # ========================================================================
    # SCHEMA COMPARISON
    # ========================================================================
    add_heading(doc, '5. Schema Comparison Guide', 1)
    
    add_heading(doc, 'What is Schema Comparison?', 2)
    add_paragraph(doc, 'Schema comparison analyzes the structure of tables (entities) between two environments, comparing field definitions, data types, lengths, and other metadata properties.')
    
    add_heading(doc, 'When to Use', 2)
    add_bullet_point(doc, 'Before deploying solutions to validate customizations')
    add_bullet_point(doc, 'After solution import to verify all fields were created')
    add_bullet_point(doc, 'Troubleshooting missing or incorrect field configurations')
    add_bullet_point(doc, 'Documentation of environment differences')
    
    add_heading(doc, 'Step-by-Step Instructions', 2)
    
    add_numbered_list(doc, 'Select Schema Comparison from the main menu (Option 1)')
    add_numbered_list(doc, 'Enter the table logical name (e.g., contact, account, mash_customtable)')
    add_numbered_list(doc, 'Wait for the tool to fetch metadata from both environments')
    add_numbered_list(doc, 'Review the console summary showing:')
    add_bullet_point(doc, 'Fields only in source', level=1)
    add_bullet_point(doc, 'Fields only in target', level=1)
    add_bullet_point(doc, 'Fields with differences', level=1)
    add_bullet_point(doc, 'Matching fields', level=1)
    add_numbered_list(doc, 'Choose to generate an Excel report (y/n)')
    add_numbered_list(doc, 'Excel file is auto-named: schema_comparison_{tablename}_{timestamp}.xlsx')
    
    add_heading(doc, 'Finding Table Logical Names', 2)
    add_paragraph(doc, 'To find a table\'s logical name in D365:')
    add_numbered_list(doc, 'Go to Settings > Customizations > Customize the System')
    add_numbered_list(doc, 'Expand Entities')
    add_numbered_list(doc, 'Look at the "Name" column (not "Display Name")')
    add_paragraph(doc, 'Common examples:')
    add_bullet_point(doc, 'contact - Contact table')
    add_bullet_point(doc, 'account - Account table')
    add_bullet_point(doc, 'incident - Case table')
    add_bullet_point(doc, 'mash_customtable - Custom table with prefix')
    
    add_heading(doc, 'Excel Report Structure', 2)
    add_paragraph(doc, 'The schema comparison Excel report contains:')
    add_bullet_point(doc, 'Summary Sheet: Overview with statistics')
    add_bullet_point(doc, 'Only in Source: Fields missing in target environment')
    add_bullet_point(doc, 'Only in Target: New fields in target environment')
    add_bullet_point(doc, 'Field Differences: Properties that differ between environments')
    add_bullet_point(doc, 'Matching Fields: Fields that are identical')
    
    doc.add_page_break()
    
    # ========================================================================
    # DATA COMPARISON
    # ========================================================================
    add_heading(doc, '6. Data Comparison Guide', 1)
    
    add_heading(doc, 'What is Data Comparison?', 2)
    add_paragraph(doc, 'Data comparison analyzes actual records between environments, identifying missing records, field value differences, and comparing related child records through One-to-Many relationships.')
    
    add_heading(doc, 'Key Features', 2)
    add_bullet_point(doc, 'GUID-based matching: Compares records by their unique identifiers')
    add_bullet_point(doc, 'System field exclusion: Automatically excludes 35+ system fields')
    add_bullet_point(doc, 'Relationship support: Compares related records (subgrids)')
    add_bullet_point(doc, 'GUID mismatch detection: Identifies lookup field discrepancies')
    add_bullet_point(doc, 'Field-level differences: Shows exactly which fields differ')
    
    add_heading(doc, 'Step-by-Step Instructions', 2)
    
    add_numbered_list(doc, 'Select Data Comparison from the main menu (Option 2)')
    add_numbered_list(doc, 'Enter the table logical name')
    add_numbered_list(doc, 'Choose comparison scope:')
    add_bullet_point(doc, 'Option 1: Main records only (faster)', level=1)
    add_bullet_point(doc, 'Option 2: Main + related records (comprehensive)', level=1)
    add_numbered_list(doc, 'Wait for data fetching and comparison')
    add_numbered_list(doc, 'Review the detailed console summary')
    add_numbered_list(doc, 'Generate Excel report when prompted')
    
    add_heading(doc, 'System Fields Automatically Excluded', 2)
    add_paragraph(doc, 'The tool excludes 35 system fields from comparison:')
    add_bullet_point(doc, 'Timestamps: modifiedon, createdon, overriddencreatedon')
    add_bullet_point(doc, 'Ownership: ownerid, owninguser, owningteam, owningbusinessunit')
    add_bullet_point(doc, 'System: versionnumber, importsequencenumber, timezoneruleversionnumber')
    add_bullet_point(doc, 'Created/Modified by: createdby, modifiedby, createdonbehalfby, modifiedonbehalfby')
    add_bullet_point(doc, 'Workflow: processid, stageid, traversedpath')
    add_paragraph(doc, 'This focuses the comparison on business-relevant data.')
    
    add_heading(doc, 'Excel Report Structure', 2)
    add_paragraph(doc, 'Data comparison reports include:')
    add_bullet_point(doc, 'Summary: Record counts and statistics')
    add_bullet_point(doc, 'Only in Source: Records missing in target')
    add_bullet_point(doc, 'Only in Target: Extra records in target')
    add_bullet_point(doc, 'Field Mismatches: Records with different field values')
    add_bullet_point(doc, 'GUID Mismatches: Lookup field discrepancies')
    add_bullet_point(doc, 'Matching Records: Identical records')
    add_bullet_point(doc, 'Child Entity Sheets: Related record comparisons')
    
    doc.add_page_break()
    
    # ========================================================================
    # FLOW COMPARISON
    # ========================================================================
    add_heading(doc, '7. Flow Comparison Guide', 1)
    
    add_heading(doc, 'What is Flow Comparison?', 2)
    add_paragraph(doc, 'Flow comparison analyzes Power Automate flow definitions between environments, using SHA256 hashing and JSON normalization to detect meaningful differences while ignoring environment-specific GUIDs and timestamps.')
    
    add_heading(doc, 'Key Features', 2)
    add_bullet_point(doc, 'SHA256 hashing: Efficiently detects identical vs different flows')
    add_bullet_point(doc, 'JSON normalization: Removes environment-specific details')
    add_bullet_point(doc, 'Action-level diff: Shows exactly which actions differ')
    add_bullet_point(doc, 'Connection masking: Ignores connection reference GUIDs')
    add_bullet_point(doc, 'Missing flow detection: Identifies flows not in target')
    
    add_heading(doc, 'Step-by-Step Instructions', 2)
    
    add_numbered_list(doc, 'Select Flow Comparison from the main menu (Option 3)')
    add_numbered_list(doc, 'Enter the flow name to compare')
    add_numbered_list(doc, 'Wait for flow fetching and analysis')
    add_numbered_list(doc, 'Review comparison results:')
    add_bullet_point(doc, 'Identical flows: Same definition', level=1)
    add_bullet_point(doc, 'Different flows: Action differences detected', level=1)
    add_bullet_point(doc, 'Missing flows: Not in target environment', level=1)
    add_numbered_list(doc, 'Generate Excel report with 6 sheets')
    
    add_heading(doc, 'Excel Report Structure', 2)
    add_bullet_point(doc, 'Summary: Flow counts and overall statistics')
    add_bullet_point(doc, 'Identical Flows: Flows with matching definitions')
    add_bullet_point(doc, 'Different Flows: Flows with detected differences')
    add_bullet_point(doc, 'Action Differences: Detailed action-level changes')
    add_bullet_point(doc, 'Missing Flows: Flows not found in target')
    add_bullet_point(doc, 'Errors: Any processing errors encountered')
    
    doc.add_page_break()
    
    # ========================================================================
    # SOLUTION COMPARISON
    # ========================================================================
    add_heading(doc, '8. Solution Comparison Guide', 1)
    
    add_heading(doc, 'What is Solution Comparison?', 2)
    add_paragraph(doc, 'Solution comparison analyzes all components within a solution across environments, comparing 160+ component types including entities, forms, workflows, web resources, plugins, canvas apps, and more.')
    
    add_heading(doc, 'Key Features', 2)
    add_bullet_point(doc, '160+ component types: Comprehensive coverage')
    add_bullet_point(doc, 'Component-level granularity: Individual component comparison')
    add_bullet_point(doc, 'Type breakdown: Grouped statistics by component type')
    add_bullet_point(doc, 'Version tracking: Solution version comparison')
    add_bullet_point(doc, 'Managed status: Identifies managed vs unmanaged')
    
    add_heading(doc, 'Component Types Supported', 2)
    add_paragraph(doc, 'The tool compares all these component types:')
    add_bullet_point(doc, 'Entities, Attributes, Forms, Views, Relationships')
    add_bullet_point(doc, 'Workflows, Business Rules, Actions')
    add_bullet_point(doc, 'Web Resources, Site Maps')
    add_bullet_point(doc, 'Plugin Assemblies, Plugin Types, SDK Steps')
    add_bullet_point(doc, 'Canvas Apps, Model-driven Apps, App Modules')
    add_bullet_point(doc, 'Environment Variables, Connection References')
    add_bullet_point(doc, 'Reports, Dashboards, Charts')
    add_bullet_point(doc, 'And 150+ more types')
    
    add_heading(doc, 'Step-by-Step Instructions', 2)
    
    add_numbered_list(doc, 'Select Solution Comparison from main menu (Option 4)')
    add_numbered_list(doc, 'Enter the solution unique name (case-sensitive)')
    add_numbered_list(doc, 'Wait for component fetching from both environments')
    add_numbered_list(doc, 'Review console summary showing:')
    add_bullet_point(doc, 'Solution version and managed status', level=1)
    add_bullet_point(doc, 'Total component counts', level=1)
    add_bullet_point(doc, 'Component type breakdown', level=1)
    add_bullet_point(doc, 'Missing/extra components', level=1)
    add_numbered_list(doc, 'Generate Excel report with 5 sheets')
    
    add_heading(doc, 'Excel Report Structure', 2)
    add_bullet_point(doc, 'Summary: Solution metadata and component counts')
    add_bullet_point(doc, 'Component Type Summary: Breakdown by type with source/target/common counts')
    add_bullet_point(doc, 'Only in Source: Components missing from target (orange highlight)')
    add_bullet_point(doc, 'Only in Target: Extra components in target (blue highlight)')
    add_bullet_point(doc, 'Common Components: Components in both environments (green highlight)')
    
    add_heading(doc, 'Use Cases', 2)
    add_bullet_point(doc, 'Solution Migration Validation: Verify all components deployed')
    add_bullet_point(doc, 'Environment Parity Check: Ensure UAT matches Production')
    add_bullet_point(doc, 'Dependency Analysis: Check base solution components exist')
    add_bullet_point(doc, 'Solution Cleanup: Identify components to remove')
    add_bullet_point(doc, 'Version Control: Track changes between versions')
    
    doc.add_page_break()
    
    # ========================================================================
    # AUTHENTICATION METHODS
    # ========================================================================
    add_heading(doc, '9. Authentication Methods', 1)
    
    add_heading(doc, 'Option 1: Interactive Browser Login (Recommended)', 2)
    add_paragraph(doc, 'This method opens your browser for sign-in and works with:')
    add_bullet_point(doc, 'Multi-factor authentication (MFA)')
    add_bullet_point(doc, 'Managed devices')
    add_bullet_point(doc, 'Conditional access policies')
    add_bullet_point(doc, 'Azure AD-integrated accounts')
    
    add_paragraph(doc, 'How it works:', bold=True)
    add_numbered_list(doc, 'Tool opens browser on localhost:8765')
    add_numbered_list(doc, 'You sign in with your credentials')
    add_numbered_list(doc, 'Browser redirects back to the tool')
    add_numbered_list(doc, 'Tool captures the authentication token')
    
    add_heading(doc, 'Option 2: Service Principal (For Automation)', 2)
    add_paragraph(doc, 'Use this method for:')
    add_bullet_point(doc, 'Automated/scheduled comparisons')
    add_bullet_point(doc, 'CI/CD pipelines')
    add_bullet_point(doc, 'Server-based execution')
    add_bullet_point(doc, 'Non-interactive scenarios')
    
    add_heading(doc, 'Setting Up Service Principal', 2)
    
    add_numbered_list(doc, 'Register Application in Azure AD:')
    add_bullet_point(doc, 'Go to Azure Portal > Azure Active Directory > App Registrations', level=1)
    add_bullet_point(doc, 'Click "New registration"', level=1)
    add_bullet_point(doc, 'Name: "D365 Comparison Tool"', level=1)
    add_bullet_point(doc, 'Click "Register"', level=1)
    
    add_numbered_list(doc, 'Note Credentials:')
    add_bullet_point(doc, 'Copy Application (client) ID', level=1)
    add_bullet_point(doc, 'Copy Directory (tenant) ID', level=1)
    
    add_numbered_list(doc, 'Create Client Secret:')
    add_bullet_point(doc, 'Go to "Certificates & secrets"', level=1)
    add_bullet_point(doc, 'Click "New client secret"', level=1)
    add_bullet_point(doc, 'Copy the secret value (shown only once!)', level=1)
    
    add_numbered_list(doc, 'Set API Permissions:')
    add_bullet_point(doc, 'Go to "API permissions"', level=1)
    add_bullet_point(doc, 'Add permission > Dynamics CRM', level=1)
    add_bullet_point(doc, 'Select "user_impersonation"', level=1)
    add_bullet_point(doc, 'Grant admin consent', level=1)
    
    add_numbered_list(doc, 'Create Application User in D365:')
    add_bullet_point(doc, 'Go to Settings > Security > Users', level=1)
    add_bullet_point(doc, 'Switch to "Application Users" view', level=1)
    add_bullet_point(doc, 'Create new user with Application ID', level=1)
    add_bullet_point(doc, 'Assign System Administrator role', level=1)
    
    doc.add_page_break()
    
    # ========================================================================
    # EXCEL REPORTS
    # ========================================================================
    add_heading(doc, '10. Understanding Excel Reports', 1)
    
    add_heading(doc, 'Automatic File Naming', 2)
    add_paragraph(doc, 'Excel files are automatically named using the pattern:')
    add_code_block(doc, '{comparison_type}_{item_name}_{timestamp}.xlsx')
    add_paragraph(doc, 'Examples:')
    add_bullet_point(doc, 'schema_comparison_contact_20251123_143022.xlsx')
    add_bullet_point(doc, 'data_comparison_account_20251123_145530.xlsx')
    add_bullet_point(doc, 'flow_comparison_MyFlow_20251123_150145.xlsx')
    add_bullet_point(doc, 'solution_comparison_CrmSolution_20251123_152233.xlsx')
    
    add_heading(doc, 'Color Coding', 2)
    add_paragraph(doc, 'Reports use consistent color coding:')
    add_bullet_point(doc, 'Red/Pink: Items only in one environment (missing in other)')
    add_bullet_point(doc, 'Yellow/Orange: Differences or mismatches')
    add_bullet_point(doc, 'Green: Matching/identical items')
    add_bullet_point(doc, 'Blue: Informational (e.g., extra items in target)')
    
    add_heading(doc, 'Environment Names in Reports', 2)
    add_paragraph(doc, 'Reports use actual environment names instead of generic "Source/Target":')
    add_bullet_point(doc, 'Sheet names: "Only in mashtest" instead of "Only in Source"')
    add_bullet_point(doc, 'Column headers: "mashtest Count" instead of "Source Count"')
    add_bullet_point(doc, 'Labels: Environment-specific throughout')
    
    add_heading(doc, 'Common Sheet Types', 2)
    
    add_paragraph(doc, 'Summary Sheet', bold=True)
    add_bullet_point(doc, 'Overview statistics and counts')
    add_bullet_point(doc, 'Environment details and comparison metadata')
    add_bullet_point(doc, 'Generation timestamp')
    
    add_paragraph(doc, 'Difference Sheets', bold=True)
    add_bullet_point(doc, 'Items only in source environment')
    add_bullet_point(doc, 'Items only in target environment')
    add_bullet_point(doc, 'Items with property or value differences')
    
    add_paragraph(doc, 'Matching Sheet', bold=True)
    add_bullet_point(doc, 'Items that are identical')
    add_bullet_point(doc, 'Provides confirmation of parity')
    
    doc.add_page_break()
    
    # ========================================================================
    # TROUBLESHOOTING
    # ========================================================================
    add_heading(doc, '11. Troubleshooting', 1)
    
    add_heading(doc, 'Authentication Issues', 2)
    
    add_paragraph(doc, 'Problem: "Authentication failed" or "401 Unauthorized"', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Verify credentials are correct')
    add_bullet_point(doc, 'Check user has access to both environments')
    add_bullet_point(doc, 'For service principal, verify API permissions granted')
    add_bullet_point(doc, 'Ensure application user created in D365')
    add_bullet_point(doc, 'Check if MFA is interfering (use browser login)')
    
    add_paragraph(doc, 'Problem: Browser doesn\'t open for interactive login', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Check if port 8765 is available')
    add_bullet_point(doc, 'Temporarily disable firewall/antivirus')
    add_bullet_point(doc, 'Try using service principal instead')
    add_bullet_point(doc, 'Run as administrator')
    
    add_heading(doc, 'API and Data Issues', 2)
    
    add_paragraph(doc, 'Problem: "Failed to fetch metadata"', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Verify environment URL is correct')
    add_bullet_point(doc, 'Check network connectivity')
    add_bullet_point(doc, 'Ensure table logical name is correct (not display name)')
    add_bullet_point(doc, 'Try accessing the URL in browser')
    add_bullet_point(doc, 'Check for VPN requirements')
    
    add_paragraph(doc, 'Problem: "Table not found" error', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Verify using logical name, not display name')
    add_bullet_point(doc, 'Check spelling and case sensitivity')
    add_bullet_point(doc, 'Ensure table exists in both environments')
    add_bullet_point(doc, 'Check user has read permissions on the table')
    
    add_heading(doc, 'Permission Issues', 2)
    
    add_paragraph(doc, 'Problem: "Access denied" errors', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Verify user has System Administrator or System Customizer role')
    add_bullet_point(doc, 'Check read permissions on specific entities')
    add_bullet_point(doc, 'Ensure user can access Web API')
    add_bullet_point(doc, 'For application user, verify security role assignment')
    
    add_heading(doc, 'Installation Issues', 2)
    
    add_paragraph(doc, 'Problem: "No module named \'requests\'" or similar', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Re-run setup.bat')
    add_bullet_point(doc, 'Manually install: pip install -r requirements.txt')
    add_bullet_point(doc, 'Check Python version (must be 3.8+)')
    add_bullet_point(doc, 'Ensure pip is installed and working')
    
    add_paragraph(doc, 'Problem: Python not recognized', bold=True)
    add_paragraph(doc, 'Solutions:')
    add_bullet_point(doc, 'Install Python from python.org')
    add_bullet_point(doc, 'Add Python to PATH environment variable')
    add_bullet_point(doc, 'Restart command prompt after installation')
    add_bullet_point(doc, 'Use full path: C:\\Python39\\python.exe main.py')
    
    doc.add_page_break()
    
    # ========================================================================
    # TIPS & BEST PRACTICES
    # ========================================================================
    add_heading(doc, '12. Tips & Best Practices', 1)
    
    add_heading(doc, 'General Tips', 2)
    add_bullet_point(doc, 'Run comparisons before and after deployments')
    add_bullet_point(doc, 'Save Excel reports for audit trails and documentation')
    add_bullet_point(doc, 'Use descriptive filenames including dates')
    add_bullet_point(doc, 'Test on non-production environments first')
    add_bullet_point(doc, 'Keep the tool updated to latest version')
    
    add_heading(doc, 'Performance Tips', 2)
    add_bullet_point(doc, 'For large tables, consider filtering or using data comparison on smaller subsets')
    add_bullet_point(doc, 'Related record comparison takes longer - use only when needed')
    add_bullet_point(doc, 'Run during off-peak hours for large comparisons')
    add_bullet_point(doc, 'Use service principal for faster authentication')
    
    add_heading(doc, 'Data Comparison Best Practices', 2)
    add_bullet_point(doc, 'Always use GUID-based comparison (automatic)')
    add_bullet_point(doc, 'Review system field exclusions - they are excluded for good reason')
    add_bullet_point(doc, 'For lookup fields, pay attention to GUID mismatches')
    add_bullet_point(doc, 'Use related records comparison for parent-child validation')
    add_bullet_point(doc, 'Focus on business fields, not system metadata')
    
    add_heading(doc, 'Solution Comparison Best Practices', 2)
    add_bullet_point(doc, 'Use solution unique name (case-sensitive)')
    add_bullet_point(doc, 'Compare same versions for accurate results')
    add_bullet_point(doc, 'Note managed vs unmanaged status differences')
    add_bullet_point(doc, 'Investigate missing components before deployment')
    add_bullet_point(doc, 'Keep component type summary for documentation')
    
    add_heading(doc, 'Excel Report Tips', 2)
    add_bullet_point(doc, 'Review Summary sheet first for overview')
    add_bullet_point(doc, 'Filter and sort columns as needed')
    add_bullet_point(doc, 'Use Excel\'s conditional formatting for additional insights')
    add_bullet_point(doc, 'Export to PDF for sharing with non-Excel users')
    add_bullet_point(doc, 'Create a library of comparison reports for history')
    
    add_heading(doc, 'Finding Logical Names', 2)
    add_paragraph(doc, 'Tables:', bold=True)
    add_bullet_point(doc, 'Settings > Customizations > Entities > "Name" column')
    
    add_paragraph(doc, 'Fields:', bold=True)
    add_bullet_point(doc, 'Open entity customization > Fields > "Name" column')
    
    add_paragraph(doc, 'Common Patterns:', bold=True)
    add_bullet_point(doc, 'Standard tables: lowercase (contact, account, lead)')
    add_bullet_point(doc, 'Custom tables: prefix_name (mash_customtable, new_invoice)')
    add_bullet_point(doc, 'Activities: lowercase (email, phonecall, task, appointment)')
    
    add_heading(doc, 'Deployment Workflow', 2)
    add_numbered_list(doc, 'Run schema comparison BEFORE solution export')
    add_numbered_list(doc, 'Document any environment-specific customizations')
    add_numbered_list(doc, 'Export solution from source environment')
    add_numbered_list(doc, 'Import solution to target environment')
    add_numbered_list(doc, 'Run schema comparison AFTER import to validate')
    add_numbered_list(doc, 'Run data comparison if data was migrated')
    add_numbered_list(doc, 'Generate and archive all comparison reports')
    
    doc.add_page_break()
    
    # ========================================================================
    # APPENDIX
    # ========================================================================
    add_heading(doc, '13. Appendix', 1)
    
    add_heading(doc, 'File Structure', 2)
    add_code_block(doc, '''D365PythonComparison/
├── main.py                  # Entry point
├── run.bat                  # Quick launcher
├── setup.bat                # Dependency installer
├── requirements.txt         # Python packages
├── src/
│   ├── auth_manager.py     # Authentication
│   ├── schema_comparison.py # Schema logic
│   ├── data_comparison.py  # Data logic
│   ├── flow_comparison.py  # Flow logic
│   ├── solution_comparison.py # Solution logic
│   └── excel_generator.py  # Report generation
└── docs/                    # Documentation''')
    
    add_heading(doc, 'System Fields Excluded from Data Comparison', 2)
    add_paragraph(doc, 'Complete list of 35 excluded fields:')
    fields = [
        'modifiedon', 'createdon', 'overriddencreatedon', 'modifiedby', 'createdby',
        'ownerid', 'owninguser', 'owningteam', 'owningbusinessunit', 'versionnumber',
        'importsequencenumber', 'timezoneruleversionnumber', 'utcconversiontimezonecode',
        'createdonbehalfby', 'modifiedonbehalfby', '_createdby_value', '_modifiedby_value',
        '_ownerid_value', '_owninguser_value', '_owningteam_value', '_owningbusinessunit_value',
        '_createdonbehalfby_value', '_modifiedonbehalfby_value', 'processid', 'stageid',
        'traversedpath', 'statecode', 'statuscode', 'transactioncurrencyid',
        '_transactioncurrencyid_value', 'exchangerate', 'solutionid', 'supportingsolutionid',
        'componentidunique', 'componentstate'
    ]
    for field in fields:
        add_bullet_point(doc, field)
    
    add_heading(doc, 'Keyboard Shortcuts', 2)
    add_bullet_point(doc, 'Ctrl+C: Cancel operation and exit')
    add_bullet_point(doc, 'Ctrl+Z: Not supported (tool is interactive)')
    add_bullet_point(doc, 'Enter: Confirm selections')
    
    add_heading(doc, 'Command Line Options', 2)
    add_paragraph(doc, 'Currently, the tool is interactive-only. Future versions may support:')
    add_bullet_point(doc, '--source <url>: Specify source environment')
    add_bullet_point(doc, '--target <url>: Specify target environment')
    add_bullet_point(doc, '--table <name>: Specify table name')
    add_bullet_point(doc, '--type <schema|data|flow|solution>: Comparison type')
    add_bullet_point(doc, '--output <file>: Specify output filename')
    
    add_heading(doc, 'Version History', 2)
    add_bullet_point(doc, 'Version 1.0 (November 2025):')
    add_bullet_point(doc, 'Initial release', level=1)
    add_bullet_point(doc, 'Schema comparison', level=1)
    add_bullet_point(doc, 'Data comparison with relationships', level=1)
    add_bullet_point(doc, 'Flow comparison', level=1)
    add_bullet_point(doc, 'Solution comparison', level=1)
    add_bullet_point(doc, 'OAuth 2.0 authentication', level=1)
    add_bullet_point(doc, 'Excel report generation', level=1)
    
    add_heading(doc, 'Contact & Support', 2)
    add_paragraph(doc, 'For questions or issues:')
    add_bullet_point(doc, 'Check documentation in docs/ folder')
    add_bullet_point(doc, 'Review troubleshooting section')
    add_bullet_point(doc, 'Contact development team')
    add_bullet_point(doc, 'GitHub repository: CodeSpace123')
    
    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================
    output_path = 'D365_Comparison_Tool_How_To_Guide.docx'
    doc.save(output_path)
    print(f"\n✓ Word document created successfully: {output_path}")
    print(f"  Location: {output_path}")
    print(f"  Open it with Microsoft Word to view the complete guide.")
    
    return output_path

if __name__ == "__main__":
    print("Creating How-To Word document...")
    print("-" * 70)
    create_how_to_document()
    print("-" * 70)
    print("Done!")
